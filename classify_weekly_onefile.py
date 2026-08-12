# -*- coding: utf-8 -*-
"""Prepare Codex classification input and export the classified weekly report.

Classification is intentionally performed outside this Python process. ``prepare``
writes a bounded JSON payload plus a JSON Schema. GitHub Actions passes those files
to Codex. ``export`` validates the returned JSON, optionally translates text with
Google Translate, and creates the Excel/Word deliverables.

No Gemini client, endpoint, model, or API key is used by this module.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import random
import re
import time
from pathlib import Path
from typing import Any, Sequence

import pandas as pd
from classification_io import (
    build_codex_output_schema,
    build_codex_payload,
    load_codex_label_batches,
    load_classification_rules,
    write_codex_files,
)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from excel_output_utils import format_literature_worksheet


DEFAULT_CLASSIFICATION_FILE = "classification.txt"
DEFAULT_CODEX_INPUT = "output/codex/classification-input.json"
DEFAULT_CODEX_SCHEMA = "output/codex/classification-schema.json"
DEFAULT_CODEX_RESULT = "output/codex/classification-result.json"
MAX_CLASSIFICATION_CHARS = int(os.getenv("MAX_CLASSIFICATION_CHARS", "5000"))
DEFAULT_CODEX_BATCH_SIZE = int(os.getenv("CODEX_CLASSIFICATION_BATCH_SIZE", "150"))
MAX_TRANSLATION_CHARS = int(os.getenv("MAX_ABSTRACT_CHARS_TO_TRANSLATE", "1600"))
GOOGLE_TRANSLATE_MAX_RETRIES = int(os.getenv("GOOGLE_TRANSLATE_MAX_RETRIES", "4"))

BASE_COLUMNS = [
    "title",
    "link",
    "published",
    "source",
    "pub_date",
    "doi",
    "last_author",
    "abstract",
    "abstract_source",
    "must_have_abstract",
    "title_zh",
    "abstract_zh",
    "categories",
]


def normalize_cell(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def ensure_base_columns(frame: pd.DataFrame) -> pd.DataFrame:
    frame = frame.copy()
    for column in BASE_COLUMNS:
        if column not in frame.columns:
            frame[column] = ""
        frame[column] = frame[column].fillna("").astype(str)
    return frame


def stable_id(row: dict[str, Any]) -> str:
    link = normalize_cell(row.get("link"))
    fallback = "||".join(
        [
            normalize_cell(row.get("title")),
            normalize_cell(row.get("source")),
            normalize_cell(row.get("pub_date") or row.get("published")),
        ]
    )
    return hashlib.sha1((link or fallback).encode("utf-8", errors="ignore")).hexdigest()[:16]


def ensure_stable_ids(frame: pd.DataFrame) -> pd.DataFrame:
    frame = frame.copy()
    frame["stable_id"] = [stable_id(row) for row in frame.to_dict(orient="records")]
    if frame["stable_id"].duplicated().any():
        duplicates = frame.loc[frame["stable_id"].duplicated(), "stable_id"].tolist()
        raise ValueError(f"Input contains duplicate stable ids: {duplicates[:5]}")
    return frame


def extract_date_from_name(filename: str):
    match = re.search(r"(\d{4})[-_]?(\d{2})[-_]?(\d{2})", Path(filename).stem)
    if not match:
        return None
    try:
        return tuple(map(int, match.groups()))
    except ValueError:
        return None


def pick_latest_weekly_csv(folder: str | Path = "output/weekly") -> Path:
    root = Path(folder)
    candidates = [
        path
        for path in root.glob("weekly_news_with_abstract_*.csv")
        if "_translated" not in path.stem
    ]
    if not candidates:
        raise FileNotFoundError(f"No weekly CSV found in {root.resolve()}")
    return max(candidates, key=lambda path: (extract_date_from_name(path.name) or (0, 0, 0), path.stat().st_mtime))


def load_weekly_csv(path: str | Path) -> pd.DataFrame:
    return ensure_stable_ids(
        ensure_base_columns(
            pd.read_csv(path, encoding="utf-8-sig", keep_default_na=False)
        )
    )


def bounded_text(value: Any, limit: int = MAX_CLASSIFICATION_CHARS) -> str:
    text = re.sub(r"\s+", " ", normalize_cell(value)).strip()
    if len(text) <= limit:
        return text
    half = limit // 2
    return f"{text[:half]} ... {text[-half:]}"


def prepare_codex_classification(
    frame: pd.DataFrame,
    rules_path: str | Path,
    input_path: str | Path,
    schema_path: str | Path,
    batch_size: int = DEFAULT_CODEX_BATCH_SIZE,
) -> list[Path]:
    if batch_size < 1:
        raise ValueError("Codex classification batch size must be positive")
    rules = load_classification_rules(rules_path)
    items = [
        {
            "id": row["stable_id"],
            "title": bounded_text(row.get("title"), 1000),
            "abstract": bounded_text(row.get("abstract")),
        }
        for row in frame.to_dict(orient="records")
    ]
    if not items:
        raise ValueError("Weekly CSV contains no records to classify")
    base_input_path = Path(input_path)
    schema = build_codex_output_schema(rules)
    batches: list[Path] = []
    for batch_number, start in enumerate(range(0, len(items), batch_size)):
        batch_path = base_input_path.with_name(
            f"{base_input_path.stem}-{batch_number:03d}{base_input_path.suffix}"
        )
        write_codex_files(
            build_codex_payload(items[start : start + batch_size], rules),
            schema,
            batch_path,
            schema_path,
        )
        batches.append(batch_path)

    manifest_path = base_input_path.with_name("classification-batches.json")
    manifest_path.write_text(
        json.dumps([f"{number:03d}" for number in range(len(batches))]),
        encoding="utf-8",
    )
    return batches


def _atomic_write_csv(frame: pd.DataFrame, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_name(path.name + ".tmp")
    frame.to_csv(temporary, index=False, encoding="utf-8-sig")
    temporary.replace(path)


def _load_translation_checkpoint(frame: pd.DataFrame, path: Path) -> pd.DataFrame:
    if not path.exists():
        return frame
    cached = pd.read_csv(path, encoding="utf-8-sig", keep_default_na=False)
    if "stable_id" not in cached.columns:
        return frame
    cached = cached.drop_duplicates("stable_id", keep="last").set_index("stable_id")
    for column in ("title_zh", "abstract_zh"):
        if column not in cached.columns:
            continue
        for index, item_id in frame["stable_id"].items():
            if not normalize_cell(frame.at[index, column]) and item_id in cached.index:
                frame.at[index, column] = normalize_cell(cached.at[item_id, column])
    return frame


def google_translate_texts(texts: Sequence[str], label: str) -> list[str]:
    """Translate through the public Google Translate endpoint; never Gemini."""
    import requests

    translated: list[str] = []
    for position, raw_text in enumerate(texts, start=1):
        text = normalize_cell(raw_text)
        if not text:
            translated.append("")
            continue
        last_error: Exception | None = None
        for attempt in range(1, GOOGLE_TRANSLATE_MAX_RETRIES + 1):
            try:
                response = requests.get(
                    "https://translate.googleapis.com/translate_a/single",
                    params={"client": "gtx", "sl": "auto", "tl": "zh-CN", "dt": "t", "q": text},
                    timeout=45,
                )
                response.raise_for_status()
                data = response.json()
                pieces = data[0] if isinstance(data, list) and data and isinstance(data[0], list) else []
                translated.append("".join(str(piece[0]) for piece in pieces if isinstance(piece, list) and piece).strip())
                break
            except Exception as error:  # network retry boundary
                last_error = error
                if attempt < GOOGLE_TRANSLATE_MAX_RETRIES:
                    time.sleep(min(15, 2**attempt) + random.uniform(0, 1))
        else:
            raise RuntimeError(f"Google Translate failed for {label} item {position}: {last_error}")
    return translated


def enrich_translation(
    frame: pd.DataFrame,
    provider: str,
    checkpoint_path: Path | None,
) -> pd.DataFrame:
    provider = provider.strip().lower()
    if provider == "none":
        return frame
    if provider != "google":
        raise ValueError("Translation provider must be google or none")

    frame = frame.copy()
    if checkpoint_path:
        frame = _load_translation_checkpoint(frame, checkpoint_path)

    for source_column, target_column, limit in (
        ("title", "title_zh", 1000),
        ("abstract", "abstract_zh", MAX_TRANSLATION_CHARS),
    ):
        pending = frame.index[
            (frame[source_column].str.strip() != "")
            & (frame[target_column].str.strip() == "")
        ].tolist()
        for start in range(0, len(pending), 10):
            indices = pending[start : start + 10]
            texts = [normalize_cell(frame.at[index, source_column])[:limit] for index in indices]
            values = google_translate_texts(texts, target_column)
            for index, value in zip(indices, values):
                frame.at[index, target_column] = value
            if checkpoint_path:
                _atomic_write_csv(
                    frame[["stable_id", "title_zh", "abstract_zh"]],
                    checkpoint_path,
                )
    return frame


def apply_codex_result(
    frame: pd.DataFrame,
    rules_path: str | Path,
    result_path: str | Path,
) -> tuple[pd.DataFrame, list[list[str]], list[str]]:
    rules = load_classification_rules(rules_path)
    ordered_categories = [rule.name for rule in rules]
    result_root = Path(result_path)
    result_paths = (
        sorted(result_root.glob("classification-result-*.json"))
        if result_root.is_dir()
        else [result_root]
    )
    labels_by_id = load_codex_label_batches(
        result_paths,
        frame["stable_id"].tolist(),
        ordered_categories,
    )
    labels = [labels_by_id[item_id] for item_id in frame["stable_id"]]
    frame = frame.copy()
    frame["categories"] = [";".join(values) for values in labels]
    return frame, labels, ordered_categories


def _safe_sheet_name(name: str, used: set[str]) -> str:
    base = re.sub(r"[\\/*?:\[\]]", "_", normalize_cell(name))[:31] or "Sheet"
    candidate = base
    suffix = 2
    while candidate in used:
        tail = f"_{suffix}"
        candidate = base[: 31 - len(tail)] + tail
        suffix += 1
    used.add(candidate)
    return candidate


def write_grouped_xlsx(
    frame: pd.DataFrame,
    labels: Sequence[Sequence[str]],
    categories: Sequence[str],
    output_path: str | Path,
) -> None:
    used: set[str] = set()
    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        wrote_sheet = False
        for category in categories:
            indices = [index for index, row_labels in enumerate(labels) if category in row_labels]
            if not indices:
                continue
            frame.iloc[indices].drop(columns=["stable_id"], errors="ignore").to_excel(
                writer,
                sheet_name=_safe_sheet_name(category, used),
                index=False,
            )
            format_literature_worksheet(writer.book.worksheets[-1])
            wrote_sheet = True

        uncategorized = [index for index, row_labels in enumerate(labels) if not row_labels]
        if uncategorized or not wrote_sheet:
            frame.iloc[uncategorized].drop(columns=["stable_id"], errors="ignore").to_excel(
                writer,
                sheet_name=_safe_sheet_name("未匹配", used),
                index=False,
            )
            format_literature_worksheet(writer.book.worksheets[-1])


def add_hyperlink(paragraph, url: str, text: str) -> None:
    relation_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    run.append(properties)
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _record_value(row: dict[str, Any], primary: str, fallback: str = "") -> str:
    return normalize_cell(row.get(primary)) or normalize_cell(row.get(fallback))


def write_grouped_docx(
    frame: pd.DataFrame,
    labels: Sequence[Sequence[str]],
    categories: Sequence[str],
    output_path: str | Path,
    report_title: str,
) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(report_title)
    run.bold = True
    run.font.size = Pt(18)

    records = frame.to_dict(orient="records")
    groups = [(category, [i for i, row_labels in enumerate(labels) if category in row_labels]) for category in categories]
    groups.append(("未匹配", [i for i, row_labels in enumerate(labels) if not row_labels]))
    for category, indices in groups:
        if not indices:
            continue
        heading = document.add_heading(category, level=1)
        heading.paragraph_format.space_before = Pt(10)
        for number, index in enumerate(indices, start=1):
            row = records[index]
            paragraph = document.add_paragraph()
            title_text = _record_value(row, "title_zh", "title") or "(untitled)"
            title_run = paragraph.add_run(f"{number}. {title_text}")
            title_run.bold = True
            title_run.font.size = Pt(11)
            original_title = normalize_cell(row.get("title"))
            if original_title and original_title != title_text:
                document.add_paragraph(original_title)
            meta = " | ".join(
                value
                for value in (
                    normalize_cell(row.get("source")),
                    _record_value(row, "pub_date", "published"),
                    normalize_cell(row.get("last_author")),
                )
                if value
            )
            if meta:
                document.add_paragraph(meta)
            abstract = _record_value(row, "abstract_zh", "abstract")
            if abstract:
                document.add_paragraph(abstract)
            link = normalize_cell(row.get("link"))
            doi = normalize_cell(row.get("doi"))
            target = link or (f"https://doi.org/{doi}" if doi.startswith("10.") else "")
            if target:
                link_paragraph = document.add_paragraph()
                add_hyperlink(link_paragraph, target, target)

    document.save(output_path)


def resolve_input(value: str) -> Path:
    path = Path(value) if value else pick_latest_weekly_csv()
    if not path.exists():
        raise FileNotFoundError(path)
    return path


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("mode", choices=["prepare", "export"])
    parser.add_argument("-i", "--input", default="")
    parser.add_argument("-c", "--classification", default=DEFAULT_CLASSIFICATION_FILE)
    parser.add_argument("--codex-input", default=DEFAULT_CODEX_INPUT)
    parser.add_argument("--codex-schema", default=DEFAULT_CODEX_SCHEMA)
    parser.add_argument("--codex-result", default=DEFAULT_CODEX_RESULT)
    parser.add_argument("--codex-batch-size", type=int, default=DEFAULT_CODEX_BATCH_SIZE)
    parser.add_argument("--translation-provider", choices=["google", "none"], default="google")
    parser.add_argument("--output-suffix", default="_translated")
    parser.add_argument("--save-translated-csv", action="store_true")
    parser.add_argument("--checkpoint-dir", default="output/checkpoints")
    parser.add_argument("--no-resume", action="store_true")
    parser.add_argument("--report-title", default="Tech Tracking Digest")
    args = parser.parse_args()

    input_path = resolve_input(args.input)
    frame = load_weekly_csv(input_path)
    print(f"[io] input={input_path} rows={len(frame)}", flush=True)

    if args.mode == "prepare":
        batches = prepare_codex_classification(
            frame,
            args.classification,
            args.codex_input,
            args.codex_schema,
            args.codex_batch_size,
        )
        print(f"[io] Codex batches: {len(batches)}", flush=True)
        print(f"[io] Codex schema: {args.codex_schema}", flush=True)
        return

    suffix = args.output_suffix if args.output_suffix.startswith("_") else f"_{args.output_suffix}"
    checkpoint = None
    if not args.no_resume:
        checkpoint = Path(args.checkpoint_dir) / f"{input_path.stem}{suffix}_translation.csv"
    frame = enrich_translation(frame, args.translation_provider, checkpoint)
    frame, labels, categories = apply_codex_result(
        frame,
        args.classification,
        args.codex_result,
    )

    if args.save_translated_csv:
        translated_csv = input_path.with_name(f"{input_path.stem}{suffix}.csv")
        _atomic_write_csv(frame.drop(columns=["stable_id"], errors="ignore"), translated_csv)
        print(f"[io] translated CSV: {translated_csv}", flush=True)

    xlsx_path = input_path.with_name(f"{input_path.stem}{suffix}.xlsx")
    docx_path = xlsx_path.with_suffix(".docx")
    write_grouped_xlsx(frame, labels, categories, xlsx_path)
    write_grouped_docx(frame, labels, categories, docx_path, args.report_title)
    print(f"[io] XLSX: {xlsx_path}", flush=True)
    print(f"[io] DOCX: {docx_path}", flush=True)


if __name__ == "__main__":
    main()
